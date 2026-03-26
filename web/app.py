"""Flask backend for handwriting recognition/synthesis demo."""

import base64
import os
from io import BytesIO
from pathlib import Path

import cv2
import numpy as np
import pytesseract
from docx import Document
from docx.shared import Inches
from flask import Flask, request, render_template, jsonify, send_file
from PIL import Image
from dotenv import load_dotenv
from pypdf import PdfReader

from src.predict import recognize_digit
from src.synthesis import generate_handwriting


PROJECT_ROOT = Path(__file__).resolve().parents[1]
load_dotenv(PROJECT_ROOT / ".env")

app = Flask(__name__)
app.config["SECRET_KEY"] = os.environ.get("FLASK_SECRET_KEY", "convo-script-dev-secret")

WINDOWS_TESSERACT_PATH = Path("C:/Program Files/Tesseract-OCR/tesseract.exe")
if os.name == "nt" and WINDOWS_TESSERACT_PATH.exists():
    pytesseract.pytesseract.tesseract_cmd = str(WINDOWS_TESSERACT_PATH)


def _collect_table_text(table) -> list[str]:
    """Collect non-empty text recursively from a DOCX table."""
    values: list[str] = []
    for row in table.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                values.append(cell_text)
            for nested_table in cell.tables:
                values.extend(_collect_table_text(nested_table))
    return values


def _extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX body, tables, and headers/footers."""
    doc = Document(BytesIO(file_bytes))
    chunks: list[str] = []

    # Main body paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            chunks.append(text)

    # Main body tables
    for table in doc.tables:
        chunks.extend(_collect_table_text(table))

    # Header/footer content (common in resumes)
    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        for table in section.header.tables:
            chunks.extend(_collect_table_text(table))

        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)
        for table in section.footer.tables:
            chunks.extend(_collect_table_text(table))

    # Remove duplicates while preserving order.
    deduped = list(dict.fromkeys(chunks))
    return "\n".join(deduped).strip()


def _extract_text_from_uploaded_file(uploaded_file) -> str:
    """Extract text content from supported uploaded file types."""
    filename = (uploaded_file.filename or "").lower()
    suffix = Path(filename).suffix
    file_bytes = uploaded_file.read()

    if suffix == ".txt":
        return file_bytes.decode("utf-8", errors="ignore").strip()

    if suffix == ".docx":
        return _extract_text_from_docx_bytes(file_bytes)

    if suffix == ".pdf":
        reader = PdfReader(BytesIO(file_bytes))
        text_chunks = []
        for page in reader.pages:
            text_chunks.append(page.extract_text() or "")
        return "\n".join(text_chunks).strip()

    raise ValueError("Unsupported file type. Upload .txt, .docx, or .pdf")


def _extract_text_from_handwriting_image(image_bytes: bytes) -> str:
    """Extract digital text from a handwritten image using OCR."""
    image_array = np.frombuffer(image_bytes, dtype=np.uint8)
    image = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
    if image is None:
        raise ValueError("Invalid image file")

    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    denoised = cv2.GaussianBlur(gray, (3, 3), 0)
    binary = cv2.adaptiveThreshold(
        denoised,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        31,
        11,
    )

    pil_image = Image.fromarray(binary)
    try:
        text = pytesseract.image_to_string(pil_image, config="--oem 3 --psm 6")
    except pytesseract.pytesseract.TesseractNotFoundError as exc:
        raise RuntimeError(
            "Tesseract OCR is not installed on this machine. Install Tesseract and try again."
        ) from exc

    return text.strip()


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/health", methods=["GET"])
def health():
    return jsonify({"status": "ok"})


@app.route("/predict", methods=["POST"])
def predict():
    image_file = request.files.get("image")
    if image_file is None:
        return jsonify({"error": "No image file provided"}), 400

    data = np.frombuffer(image_file.read(), dtype=np.uint8)
    image = cv2.imdecode(data, cv2.IMREAD_GRAYSCALE)
    if image is None:
        return jsonify({"error": "Invalid image file"}), 400

    try:
        digit, confidence = recognize_digit(image)
    except FileNotFoundError as exc:
        return jsonify({"error": str(exc)}), 404

    return jsonify({"digit": digit, "confidence": round(confidence, 4)})


@app.route("/recognize", methods=["POST"])
def recognize():
    image_file = request.files.get("image")
    if image_file is None:
        return jsonify({"error": "No image file provided"}), 400

    img = cv2.imdecode(np.frombuffer(image_file.read(), np.uint8), cv2.IMREAD_COLOR)
    if img is None:
        return jsonify({"error": "Invalid image file"}), 400

    try:
        digit, _ = recognize_digit(img)
    except FileNotFoundError as exc:
        return jsonify({"error": str(exc)}), 404

    return jsonify({"recognized_digit": digit})


@app.route("/generate", methods=["POST"])
def generate():
    payload = request.get_json(silent=True) or {}
    text = payload.get("text", "")
    paper_style = (payload.get("paperStyle") or "plain").strip().lower()
    handwriting_style = (payload.get("handwritingStyle") or "cursive").strip().lower()
    if not text:
        return jsonify({"error": "Text is required"}), 400

    image = generate_handwriting(text, paper_style=paper_style, handwriting_style=handwriting_style)
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    buffer.seek(0)

    return send_file(buffer, mimetype="image/png")


@app.route("/synthesize", methods=["POST"])
def synthesize():
    text = request.form.get("text", "")
    paper_style = request.form.get("paperStyle", "plain").strip().lower()
    handwriting_style = request.form.get("handwritingStyle", "cursive").strip().lower()
    if not text:
        return jsonify({"error": "Text is required"}), 400

    img = generate_handwriting(text, paper_style=paper_style, handwriting_style=handwriting_style)
    output_path = Path(__file__).resolve().parent / "static" / "generated.png"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    img.save(output_path)
    return send_file(output_path, mimetype="image/png")


@app.route("/convert-file", methods=["POST"])
def convert_file():
    uploaded_file = request.files.get("file")
    paper_style = request.form.get("paperStyle", "plain").strip().lower()
    handwriting_style = request.form.get("handwritingStyle", "cursive").strip().lower()
    if uploaded_file is None or not uploaded_file.filename:
        return jsonify({"error": "No file uploaded"}), 400

    try:
        text = _extract_text_from_uploaded_file(uploaded_file)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400

    if not text:
        return jsonify({"error": "Could not extract any text from this file"}), 400

    image = generate_handwriting(text, paper_style=paper_style, handwriting_style=handwriting_style)

    output_path = Path(__file__).resolve().parent / "static" / "generated.png"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    image.save(output_path)

    image_buffer = BytesIO()
    image.save(image_buffer, format="PNG")
    image_data = base64.b64encode(image_buffer.getvalue()).decode("utf-8")

    return jsonify(
        {
            "extracted_text": text,
            "image_data": f"data:image/png;base64,{image_data}",
        }
    )


@app.route("/handwriting-to-text", methods=["POST"])
def handwriting_to_text():
    image_file = request.files.get("image")
    if image_file is None or not image_file.filename:
        return jsonify({"error": "No handwritten image uploaded"}), 400

    try:
        text = _extract_text_from_handwriting_image(image_file.read())
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    except RuntimeError as exc:
        return jsonify({"error": str(exc)}), 500

    if not text:
        return jsonify({"error": "No readable text detected in the uploaded image"}), 400

    return jsonify({"extracted_text": text})


@app.route("/download-handwriting", methods=["POST"])
def download_handwriting():
    text = request.form.get("text", "").strip()
    file_format = request.form.get("format", "pdf").strip().lower()
    paper_style = request.form.get("paperStyle", "plain").strip().lower()
    handwriting_style = request.form.get("handwritingStyle", "cursive").strip().lower()

    if not text:
        return jsonify({"error": "Text is required"}), 400

    image = generate_handwriting(text, paper_style=paper_style, handwriting_style=handwriting_style)

    if file_format == "pdf":
        pdf_buffer = BytesIO()
        image.convert("RGB").save(pdf_buffer, format="PDF")
        pdf_buffer.seek(0)
        return send_file(
            pdf_buffer,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="handwriting_output.pdf",
        )

    if file_format == "docx":
        doc = Document()
        doc.add_heading("Handwriting Output", level=1)
        doc.add_paragraph("Digital text:")
        doc.add_paragraph(text)
        doc.add_paragraph("Generated handwriting:")

        image_buffer = BytesIO()
        image.save(image_buffer, format="PNG")
        image_buffer.seek(0)
        doc.add_picture(image_buffer, width=Inches(6.0))

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        return send_file(
            docx_buffer,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name="handwriting_output.docx",
        )

    return jsonify({"error": "Unsupported format. Use pdf or docx."}), 400


if __name__ == "__main__":
    app.run(debug=True)
